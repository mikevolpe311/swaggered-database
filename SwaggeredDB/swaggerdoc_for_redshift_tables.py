from sqlalchemy import create_engine, text, inspect
import pandas as pd
import getpass
import tkinter as tk
from tkinter import filedialog
import json
import os
import yaml
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

##USE THIS TO CONNECT TO REDSHIFT, POSTGRESQL, MYSQL, SQLITE, SQL SERVER AND EXPORT TABLES TO EXCEL, WORD, SWAGGER

class DBToExcel:
    def __init__(self):
        self.engine = None
    
    def load_allowed_tables(self, system_name):
        """Load allowed tables from allowed_tables.json for specified system"""
        try:
            current_dir = os.path.dirname(os.path.abspath(__file__))
            allowed_tables_file = os.path.join(current_dir, 'allowed_tables.json')
            
            if not os.path.exists(allowed_tables_file):
                print(f"allowed_tables.json not found in {current_dir}")
                return None
            
            with open(allowed_tables_file, 'r') as f:
                allowed_data = json.load(f)
            
            systems = allowed_data.get('systems', {})
            if system_name not in systems:
                print(f"System '{system_name}' not found in allowed_tables.json")
                available_systems = list(systems.keys())
                if available_systems:
                    print(f"Available systems: {', '.join(available_systems)}")
                return None
            
            return systems[system_name].get('allowed_tables', [])
            
        except Exception as e:
            print(f"Error loading allowed_tables.json: {e}")
            return None
    
    def generate_sample_data(self, col):
        """Generate sample data based on column type for empty tables"""
        col_type = str(col['type']).lower()
        col_name = col['name'].lower()
        
        if 'int' in col_type or 'serial' in col_type:
            return ['1', '2', '3']
        elif 'varchar' in col_type or 'text' in col_type or 'char' in col_type:
            if 'name' in col_name:
                return ['Sample Name 1', 'Sample Name 2', 'Sample Name 3']
            elif 'email' in col_name:
                return ['user1@example.com', 'user2@example.com', 'user3@example.com']
            elif 'id' in col_name:
                return ['ID001', 'ID002', 'ID003']
            else:
                return ['Sample Text 1', 'Sample Text 2', 'Sample Text 3']
        elif 'date' in col_type or 'timestamp' in col_type:
            return ['2024-01-01', '2024-01-02', '2024-01-03']
        elif 'bool' in col_type:
            return ['true', 'false', 'true']
        elif 'decimal' in col_type or 'numeric' in col_type or 'float' in col_type:
            return ['10.50', '25.75', '100.00']
        else:
            return ['Sample Value 1', 'Sample Value 2', 'Sample Value 3']
    
    def connect(self, db_type, host, port, database, username, password=None):
        if password is None:
            password = getpass.getpass(f"Enter password for {username}: ")
        
        # Ask user about SSL
        use_ssl = input("Enable SSL connection? (y/n): ").lower().strip() == 'y'
        ssl_mode = 'require' if use_ssl else 'disable'
        
        connection_strings = {
            'redshift': f'redshift+psycopg2://{username}:{password}@{host}:{port}/{database}',
            'postgresql': f'postgresql+psycopg2://{username}:{password}@{host}:{port}/{database}',
            'mysql': f'mysql+pymysql://{username}:{password}@{host}:{port}/{database}',
            'sqlite': f'sqlite:///{database}',
            'sqlserver': f'mssql+pyodbc://{username}:{password}@{host}:{port}/{database}?driver=ODBC+Driver+17+for+SQL+Server'
        }
        
        try:
            conn_str = connection_strings[db_type.lower()]
            
            self.engine = create_engine(conn_str, connect_args={'sslmode': ssl_mode})
            with self.engine.connect() as conn:
                conn.execute(text("SELECT 1"))
            print(f"Connected to {db_type} with SSL mode: {ssl_mode}")
            #self.export_tables_to_excel()
            return True
        except Exception as e:
            error_msg = str(e)
            if "no pg_hba.conf entry" in error_msg:
                print(f"Connection failed: Your IP address is not authorized to connect to this database.")
                print(f"Contact your database administrator to add your IP to the pg_hba.conf file.")
                print(f"Full error: {e}")
            else:
                print(f"Connection failed: {e}")
            return False
    
    def export_tables_to_excel(self, output_file=None):
        if output_file is None:
            root = tk.Tk()
            root.withdraw()
            output_file = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")],
                title="Save Excel file as"
            )
            if not output_file:
                print("Export cancelled")
                return
        inspector = inspect(self.engine)
        tables = inspector.get_table_names(schema='public')
        
        # Ask if this is test mode or subset
        test_mode = input("\nIs this test mode or subset? (y/n): ").lower().strip()
        if test_mode == 'y':
            # Ask if user wants to use allowed_tables.json
            use_allowed_tables = input("Use allowed_tables.json file for table selection? (y/n): ").lower().strip()
            if use_allowed_tables == 'y':
                system_name = input("Enter system name from allowed_tables.json: ").strip()
                requested_tables = self.load_allowed_tables(system_name)
                if requested_tables:
                    requested_tables.sort()  # Sort alphabetically
                    # Log tables that don't exist
                    for req_table in requested_tables:
                        if req_table not in tables:
                            print(f"LOG: Table '{req_table}' does not exist in database")
                    tables = [t for t in requested_tables if t in tables]  # Maintain sorted order
                    print(f"Test mode or subset: Processing {len(tables)} tables from system '{system_name}'")
                else:
                    print(f"No tables found for system '{system_name}', using first 20 tables")
                    tables = tables[:20]
            else:
                specific_tables = input("Specify table names (comma-separated) or press Enter for first 20 tables: ").strip()
                if specific_tables:
                    requested_tables = [t.strip() for t in specific_tables.split(',')]
                    # Log tables that don't exist
                    for req_table in requested_tables:
                        if req_table not in tables:
                            print(f"LOG: Table '{req_table}' does not exist in database")
                    tables = [t for t in tables if t in requested_tables]
                    print(f"Test mode or subset: Processing {len(tables)} specified tables")
                else:
                    tables = tables[:20]
                    print(f"Test mode or subset: Processing first {len(tables)} tables only")
        
        all_data = []
        
        for table in tables:
            print(f"Processing table: {table}")
            columns = inspector.get_columns(table, schema='public')
            
            # Get sample data
            sample_df = pd.DataFrame()
            table_is_empty = False
            try:
                with self.engine.connect() as conn:
                    result = conn.execute(text(f"SELECT * FROM public.{table} LIMIT 2"))
                    rows = result.fetchall()
                    if rows:
                        sample_df = pd.DataFrame(rows, columns=[col['name'] for col in columns])
                    else:
                        table_is_empty = True
                        if test_mode == 'y':
                            print(f"  LOG: Table {table} is empty (test mode or subset)")
                
                print(f"  Got {len(sample_df)} sample rows")
                
                # Export sample data to JSON and XML if available
                if not sample_df.empty:
                    base_path = os.path.dirname(output_file)
                    
                    # Limit to 2 samples for Word document
                    sample_df_limited = sample_df.head(2)
                    
                    # JSON export
                    json_file = os.path.join(base_path, f"{table}_sample.json")
                    sample_df_limited.to_json(json_file, orient='records', indent=2)
                    
                    # XML export with custom root and row names
                    xml_file = os.path.join(base_path, f"{table}_sample.xml")
                    sample_df_limited.to_xml(xml_file, index=False, root_name=f"{table}s", row_name=table)
                    
                    print(f"  Exported sample data to {table}_sample.json and {table}_sample.xml")
                    
            except Exception as e:
                print(f"  Error getting sample data: {e}")
            
            # Add to documentation - include empty tables with sample data for Swagger
            for col in columns:
                sample_values = ['', '', '']
                if not sample_df.empty and col['name'] in sample_df.columns:
                    col_data = sample_df[col['name']].fillna('NULL').astype(str).tolist()
                    for i in range(min(3, len(col_data))):
                        sample_values[i] = col_data[i]
                elif table_is_empty:
                    # Generate sample data based on column type for empty tables
                    sample_values = self.generate_sample_data(col)
                
                all_data.append({
                    'Table': table,
                    'Column': col['name'],
                    'Data_Type': str(col['type']),
                    'Mandatory': 'Y' if not col.get('nullable', True) else 'N',
                    'Default': str(col.get('default', '')),
                    'Sample_1': sample_values[0],
                    'Sample_2': sample_values[1],
                    'Sample_3': sample_values[2]
                })
        
        df = pd.DataFrame(all_data)
        df.to_excel(output_file, index=False)
        print(f"Exported {len(tables)} tables with {len(all_data)} columns to: {output_file}")
        
        # Store test_mode for use in other methods
        self.test_mode = test_mode == 'y'
        
        # Ask user if they want to create Word document
        create_word = input("\nWould you like to create a Word specification document? (y/n): ").lower().strip()
        include_xml = False
        if create_word == 'y':
            include_xml = input("Include XML samples in documentation? (y/n): ").lower().strip() == 'y'
            self.create_word_spec(tables, all_data, output_file, include_xml)
        
        # Ask user if they want to create Swagger documentation
        create_swagger = input("\nWould you like to create Swagger/OpenAPI documentation? (y/n): ").lower().strip()
        if create_swagger == 'y':
            if not create_word == 'y':
                include_xml = input("Include XML support in Swagger? (y/n): ").lower().strip() == 'y'
            self.create_swagger_spec(tables, all_data, output_file, include_xml)
    
    def create_word_spec(self, tables, all_data, excel_file, include_xml=False):
        doc = Document()
        
        # Title
        title = doc.add_heading('Database Schema Specification', 0)
        
        # Document info
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        # Count tables with sample data
        base_path = os.path.dirname(excel_file)
        tables_with_data_count = sum(1 for table in tables if os.path.exists(os.path.join(base_path, f"{table}_sample.json")))
        doc.add_paragraph(f'Total Tables with Sample Data: {tables_with_data_count}')
        if test_mode == 'y':
            doc.add_paragraph('Note: This is a test mode or subset document')
        doc.add_paragraph('')
        
        # Table of Contents - only show tables with sample data
        doc.add_heading('Table of Contents', level=1)
        tables_with_data = []
        base_path = os.path.dirname(excel_file)
        
        for table in tables:
            json_file = os.path.join(base_path, f"{table}_sample.json")
            if os.path.exists(json_file):
                tables_with_data.append(table)
        
        for i, table in enumerate(tables_with_data, 1):
            doc.add_paragraph(f'{i}. {table}', style='List Number')
        doc.add_page_break()
        
        # Table specifications - only include tables with sample data
        tables_with_data = []
        base_path = os.path.dirname(excel_file)
        
        for table in tables:
            json_file = os.path.join(base_path, f"{table}_sample.json")
            if os.path.exists(json_file):
                tables_with_data.append(table)
        
        for table in tables_with_data:
            doc.add_heading(f'Table: {table}', level=1)
            
            # Get table data
            table_data = [row for row in all_data if row['Table'] == table]
            
            if table_data:
                # API URL section
                doc.add_heading('API URL', level=2)
                
                # Create API URL table
                api_table = doc.add_table(rows=1, cols=4)
                api_table.style = 'Table Grid'
                api_hdr_cells = api_table.rows[0].cells
                
                # Format API header
                from docx.oxml.shared import qn
                from docx.oxml import parse_xml
                
                # Set API header text first
                api_hdr_cells[0].text = 'Resource'
                api_hdr_cells[1].text = 'Base URL'
                api_hdr_cells[2].text = 'Request Method'
                api_hdr_cells[3].text = 'Notes'
                
                # Then format each cell with black background
                for cell in api_hdr_cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cell.paragraphs[0].runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                    
                    # Add black background
                    shading_elm = parse_xml('<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="000000"/>')
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                
                # Add GET row
                get_row_cells = api_table.add_row().cells
                get_row_cells[0].text = table.upper()
                get_row_cells[1].text = f'BaseURL/API/{{system}}/{table}'
                get_row_cells[2].text = 'GET'
                get_row_cells[3].text = ''
                
                # Add GET ITEM row
                get_item_row_cells = api_table.add_row().cells
                get_item_row_cells[0].text = table.upper()
                get_item_row_cells[1].text = f'BaseURL/API/{{system}}/{table}/{{guid}}'
                get_item_row_cells[2].text = 'GET ITEM'
                get_item_row_cells[3].text = ''
                
                doc.add_paragraph('')
                
                # Query Parameters section
                doc.add_heading('Query Parameters', level=2)
                
                # Create parameters table
                param_table = doc.add_table(rows=1, cols=4)
                param_table.style = 'Table Grid'
                param_hdr_cells = param_table.rows[0].cells
                
                # Set parameter header text
                param_hdr_cells[0].text = 'Parameter'
                param_hdr_cells[1].text = 'Type'
                param_hdr_cells[2].text = 'Required'
                param_hdr_cells[3].text = 'Description'
                
                # Format parameter header
                for cell in param_hdr_cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cell.paragraphs[0].runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    
                    shading_elm = parse_xml('<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="4472C4"/>')
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                
                # Add parameter rows
                params = [
                    ('filter_by', 'string', 'No', 'WHERE clause condition (e.g., IMPORTEDTIME>\'2024-01-05\')'),
                    ('limit', 'integer', 'No', 'Maximum number of records to return (e.g., 300)'),
                    ('offset', 'integer', 'No', 'Number of records to skip for pagination (e.g., 0)'),
                    ('order_by', 'string', 'No', 'ORDER BY clause for sorting (e.g., IMPORTEDTIME ASC)'),
                    ('column_names', 'string', 'No', 'Comma-delimited list of column names (e.g., id,name,email)')
                ]
                
                for param_name, param_type, required, description in params:
                    param_row_cells = param_table.add_row().cells
                    param_row_cells[0].text = param_name
                    param_row_cells[1].text = param_type
                    param_row_cells[2].text = required
                    param_row_cells[3].text = description
                
                doc.add_paragraph('')
                
                # Column specifications
                doc.add_heading('Column Specifications', level=2)
                
                # Create table
                table_doc = doc.add_table(rows=1, cols=5)
                table_doc.style = 'Table Grid'
                hdr_cells = table_doc.rows[0].cells
                
                # Set header text first
                hdr_cells[0].text = 'Column Name'
                hdr_cells[1].text = 'Data Type'
                hdr_cells[2].text = 'Mandatory'
                hdr_cells[3].text = 'Description'
                hdr_cells[4].text = 'Notes / Rules'
                
                # Then format each cell
                for cell in hdr_cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cell.paragraphs[0].runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                    
                    # Add blue background
                    shading_elm = parse_xml('<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="4472C4"/>')
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                
                for col_info in table_data:
                    row_cells = table_doc.add_row().cells
                    row_cells[0].text = col_info['Column']
                    row_cells[1].text = col_info['Data_Type']
                    row_cells[2].text = col_info['Mandatory']
                    row_cells[3].text = ''
                    row_cells[4].text = ''
                
                doc.add_paragraph('')
                
                # Sample data section
                doc.add_heading('Sample Data', level=2)
                
                # Include actual JSON and XML sample data
                base_path = os.path.dirname(excel_file)
                json_file = os.path.join(base_path, f"{table}_sample.json")
                
                if os.path.exists(json_file):
                    with open(json_file, 'r') as f:
                        sample_data = json.load(f)
                    
                    doc.add_paragraph('JSON Sample:')
                    json_para = doc.add_paragraph(json.dumps(sample_data, indent=2))
                    json_para.style = 'Intense Quote'
                    
                    if include_xml:
                        xml_file = os.path.join(base_path, f"{table}_sample.xml")
                        if os.path.exists(xml_file):
                            with open(xml_file, 'r') as f:
                                xml_data = f.read()
                            
                            doc.add_paragraph('XML Sample:')
                            xml_para = doc.add_paragraph(xml_data)
                            xml_para.style = 'Intense Quote'
                    
                    doc.add_paragraph('')
            
            doc.add_page_break()
        
        # Save Word document
        word_file = excel_file.replace('.xlsx', '_specification.docx')
        doc.save(word_file)
        print(f"Word specification saved to: {word_file}")
    
    def create_swagger_spec(self, tables, all_data, excel_file, include_xml=False):
        """Generate OpenAPI/Swagger specification for all tables"""
        base_path = os.path.dirname(excel_file)
        
        # Ask user for API title and system name
        api_title = input("Enter API title for Swagger documentation (or press Enter for 'Database API'): ").strip()
        if not api_title:
            api_title = "Database API"
        
        system_name = input("Enter system name for API paths (e.g., 'australiaprod'): ").strip()
        if not system_name:
            system_name = "system"
        
        # Create main OpenAPI spec
        swagger_spec = {
            "openapi": "3.0.0",
            "info": {
                "title": api_title,
                "description": f"**DOCUMENTATION ONLY - NO LIVE DATA**\n\nThis API documentation is for reference and examples only. This documentation uses example URLs that do not connect to live data sources.\n\n---\n\nAPI documentation for Insight Database\n\n**Authentication Required:**\nAll requests must include an authorization header with an access token. A token request must be made to retrieve the access token before using these endpoints.\n\n**Example Request:**\n```\ncurl -X GET 'https://api.example.com/API/{system_name}/table_name' \\\n  -H 'authorization: YOUR_ACCESS_TOKEN' \\\n  -H 'Content-Type: application/json'\n```",
                "version": "1.0.0",
                "contact": {
                    "name": "API Support"
                }
            },
            "servers": [
                {
                    "url": "https://api.example.com/v1",
                    "description": "Production server"
                }
            ],
            "paths": {},
            "components": {
                "schemas": {},
                "parameters": {
                    "GuidParam": {
                        "name": "guid",
                        "in": "path",
                        "required": True,
                        "schema": {
                            "type": "string",
                            "format": "uuid"
                        },
                        "description": "Unique identifier"
                    }
                }
            }
        }
        
        # Process each table
        for table in tables:
            table_data = [row for row in all_data if row['Table'] == table]
            if not table_data:
                continue
            
            # Create schema for the table
            schema_properties = {}
            required_fields = []
            
            for col_info in table_data:
                col_name = col_info['Column']
                data_type = self._map_db_type_to_openapi(col_info['Data_Type'])
                
                schema_properties[col_name] = {
                    "type": data_type["type"],
                    "description": col_info['Data_Type']
                }
                
                if "format" in data_type:
                    schema_properties[col_name]["format"] = data_type["format"]
                
                # Add constraints
                if "maxLength" in data_type:
                    schema_properties[col_name]["maxLength"] = data_type["maxLength"]
                if "minimum" in data_type:
                    schema_properties[col_name]["minimum"] = data_type["minimum"]
                if "maximum" in data_type:
                    schema_properties[col_name]["maximum"] = data_type["maximum"]
                
                if col_info['Mandatory'] == 'Y':
                    required_fields.append(col_name)
                
                # Add example from sample data if available
                if col_info['Sample_1']:
                    schema_properties[col_name]["example"] = col_info['Sample_1']
            
            # Add schema to components
            swagger_spec["components"]["schemas"][table] = {
                "type": "object",
                "properties": schema_properties
            }
            
            if include_xml:
                swagger_spec["components"]["schemas"][table]["xml"] = {
                    "name": table
                }
            
            if required_fields:
                swagger_spec["components"]["schemas"][table]["required"] = required_fields
            
            # Add paths for the table
            table_path = f"/API/{system_name}/{table}"
            item_path = f"/API/{system_name}/{table}/{{guid}}"
            
            # Define common query parameters
            query_parameters = [
                {
                    "name": "filter_by",
                    "in": "query",
                    "required": False,
                    "schema": {"type": "string"},
                    "description": "WHERE clause condition (e.g., IMPORTEDTIME>'2024-01-05')",
                    "example": "IMPORTEDTIME>'2024-01-05'"
                },
                {
                    "name": "limit",
                    "in": "query",
                    "required": False,
                    "schema": {"type": "integer", "minimum": 1, "maximum": 10000},
                    "description": "Maximum number of records to return",
                    "example": 300
                },
                {
                    "name": "offset",
                    "in": "query",
                    "required": False,
                    "schema": {"type": "integer", "minimum": 0},
                    "description": "Number of records to skip for pagination",
                    "example": 0
                },
                {
                    "name": "order_by",
                    "in": "query",
                    "required": False,
                    "schema": {"type": "string"},
                    "description": "ORDER BY clause for sorting (e.g., IMPORTEDTIME ASC, name DESC)",
                    "example": "IMPORTEDTIME ASC"
                },
                {
                    "name": "column_names",
                    "in": "query",
                    "required": False,
                    "schema": {"type": "string"},
                    "description": "Comma-delimited list of column names to select (e.g., id,name,email)",
                    "example": "id,name,email"
                }
            ]
            
            # GET collection endpoint
            swagger_spec["paths"][table_path] = {
                "get": {
                    "tags": [table],
                    "summary": f"Get all {table} records",
                    "description": f"Retrieve a list of all {table} records with optional filtering, sorting, and pagination",
                    "parameters": query_parameters,
                    "responses": {
                        "200": {
                            "description": "Successful response",
                            "content": {
                                "application/json": {
                                    "schema": {
                                        "type": "object",
                                        "properties": {
                                            "data": {
                                                "type": "array",
                                                "items": {
                                                    "$ref": f"#/components/schemas/{table}"
                                                }
                                            }
                                        },
                                        "required": ["data"]
                                    }
                                }
                                # XML support commented out for now
                                # "application/xml": {
                                #     "schema": {
                                #         "type": "array",
                                #         "items": {
                                #             "$ref": f"#/components/schemas/{table}"
                                #         },
                                #         "xml": {
                                #             "name": f"{table}s",
                                #             "wrapped": True
                                #         }
                                #     }
                                # }
                            }
                        },
                        "500": {
                            "description": "Internal server error"
                        }
                    }
                }
            }
            
            # GET single item endpoint - COMMENTED OUT FOR NOW
            # swagger_spec["paths"][item_path] = {
            #     "get": {
            #         "tags": [table],
            #         "summary": f"Get {table} by ID",
            #         "description": f"Retrieve a specific {table} record by its unique identifier",
            #         "parameters": [
            #             {
            #                 "$ref": "#/components/parameters/GuidParam"
            #             }
            #         ],
            #         "responses": {
            #             "200": {
            #                 "description": "Successful response",
            #                 "content": {
            #                     "application/json": {
            #                         "schema": {
            #                             "$ref": f"#/components/schemas/{table}"
            #                         }
            #                     },
            #                     "application/xml": {
            #                         "schema": {
            #                             "$ref": f"#/components/schemas/{table}"
            #                         }
            #                     }
            #                 }
            #             },
            #             "404": {
            #                 "description": "Record not found"
            #             },
            #             "500": {
            #                 "description": "Internal server error"
            #             }
            #         }
            #     }
            # }
        
        # Save as YAML and JSON
        swagger_yaml_file = os.path.join(base_path, "api_documentation.yaml")
        swagger_json_file = os.path.join(base_path, "api_documentation.json")
        
        with open(swagger_yaml_file, 'w') as f:
            yaml.dump(swagger_spec, f, default_flow_style=False, sort_keys=False)
        
        with open(swagger_json_file, 'w') as f:
            json.dump(swagger_spec, f, indent=2)
        
        print(f"Swagger documentation saved to:")
        print(f"  YAML: {swagger_yaml_file}")
        print(f"  JSON: {swagger_json_file}")
        # Generate HTML documentation
        self.create_swagger_html(swagger_spec, base_path)
        
        print(f"\nTo view the documentation:")
        print(f"  1. Open the generated HTML file in your browser")
        print(f"  2. Or go to https://editor.swagger.io and upload the YAML/JSON file")
        print(f"  3. If HTML doesn't work, use the JSON file - it's more reliable for large specs")
    
    def create_swagger_html(self, swagger_spec, base_path):
        """Generate a standalone HTML file with Swagger UI"""
        # Use compact JSON to reduce file size
        json_spec = json.dumps(swagger_spec, separators=(',', ':'))
        
        html_content = f'''<!DOCTYPE html>
<html>
<head>
    <title>API Documentation</title>
    <link rel="stylesheet" type="text/css" href="https://unpkg.com/swagger-ui-dist@4.15.5/swagger-ui.css" />
    <style>
        html {{ box-sizing: border-box; overflow: -moz-scrollbars-vertical; overflow-y: scroll; }}
        *, *:before, *:after {{ box-sizing: inherit; }}
        body {{ margin:0; background: #fafafa; }}
    </style>
</head>
<body>
    <div id="swagger-ui"></div>
    <script src="https://unpkg.com/swagger-ui-dist@4.15.5/swagger-ui-bundle.js"></script>
    <script>
        const spec = {json_spec};
        SwaggerUIBundle({{
            spec: spec,
            dom_id: '#swagger-ui',
            deepLinking: true,
            supportedSubmitMethods: [],
            presets: [
                SwaggerUIBundle.presets.apis,
                SwaggerUIBundle.presets.standalone
            ]
        }});
    </script>
</body>
</html>'''
        
        html_file = os.path.join(base_path, "api_documentation.html")
        try:
            with open(html_file, 'w', encoding='utf-8') as f:
                f.write(html_content)
            print(f"  HTML: {html_file}")
        except Exception as e:
            print(f"  Warning: Could not create HTML file: {e}")
            print(f"  Use the JSON or YAML files with https://editor.swagger.io instead")
    
    def _map_db_type_to_openapi(self, db_type):
        """Map database types to OpenAPI types with constraints"""
        import re
        db_type_str = str(db_type)
        db_type_lower = db_type_str.lower()
        result = {}
        
        # Integer types
        if 'smallint' in db_type_lower:
            result = {"type": "integer", "minimum": -32768, "maximum": 32767}
        elif 'bigint' in db_type_lower:
            result = {"type": "integer", "format": "int64"}
        elif 'int' in db_type_lower or 'serial' in db_type_lower:
            result = {"type": "integer", "format": "int32"}
        
        # Decimal/Numeric types
        elif 'decimal' in db_type_lower or 'numeric' in db_type_lower:
            result = {"type": "number"}
            # Extract precision and scale: DECIMAL(10,2)
            match = re.search(r'\((\d+)(?:,(\d+))?\)', db_type_str)
            if match:
                precision = int(match.group(1))
                scale = int(match.group(2)) if match.group(2) else 0
                # Calculate max value based on precision and scale
                max_val = (10 ** (precision - scale)) - (10 ** -scale)
                result["maximum"] = max_val
                result["minimum"] = -max_val
        
        # Float types
        elif 'float' in db_type_lower or 'double' in db_type_lower or 'real' in db_type_lower:
            result = {"type": "number"}
        
        # Boolean
        elif 'bool' in db_type_lower:
            result = {"type": "boolean"}
        
        # Date/Time types
        elif 'timestamp' in db_type_lower or 'datetime' in db_type_lower:
            result = {"type": "string", "format": "date-time"}
        elif 'date' in db_type_lower:
            result = {"type": "string", "format": "date"}
        elif 'time' in db_type_lower:
            result = {"type": "string", "format": "time"}
        
        # UUID
        elif 'uuid' in db_type_lower:
            result = {"type": "string", "format": "uuid"}
        
        # String types with length constraints
        elif 'varchar' in db_type_lower or 'char' in db_type_lower or 'text' in db_type_lower:
            result = {"type": "string"}
            # Extract length: VARCHAR(50)
            match = re.search(r'\((\d+)\)', db_type_str)
            if match:
                length = int(match.group(1))
                result["maxLength"] = length
        
        # Default to string
        else:
            result = {"type": "string"}
        
        return result


def main():
    connector = DBToExcel()
    
    db_type = input("Database type (redshift/postgresql/mysql/sqlite/sqlserver): ")
    
    if db_type.lower() == 'sqlite':
        database = input("SQLite file path: ")
        success = connector.connect('sqlite', '', '', database, '')
    else:
        host = input("Host: ")
        port = input("Port: ")
        database = input("Database: ")
        username = input("Username: ")
        success = connector.connect(db_type, host, port, database, username)
    
    if success:
        connector.export_tables_to_excel()

if __name__ == "__main__":
    main()